"""
test_tailor.py — Dry-run unit tests for the Resume Tailor module.
=================================================================
"""

import io
import datetime
from docx import Document
from docx.shared import Pt, RGBColor

from tailor import (
    parse_docx_sections,
    get_texts,
    replace_markdown_text,
    build_tailor_prompt,
    apply_tailored_sections,
    save_doc_to_bytes,
    _parse_ai_4_sections,
    _is_section_header,
    _is_bullet,
    MODEL_CASCADE,
    DAILY_LIMIT,
)
from tailor_prompt import TAILOR_PROMPT

def _create_test_docx() -> io.BytesIO:
    doc = Document()
    doc.add_paragraph("PROFESSIONAL SUMMARY")
    doc.add_paragraph("Experienced QA engineer with 4 years of testing.")

    doc.add_paragraph("EXPERIENCE")
    doc.add_paragraph("Infosys Ltd (Client: NAB)")
    p1 = doc.add_paragraph("• Built REST Assured API frameworks achieving 90% coverage")
    if p1.runs:
        p1.runs[0].font.size = Pt(10)
        p1.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph("PROJECTS")
    doc.add_paragraph("Hybrid E2E QA")
    doc.add_paragraph("• Architected a framework")

    doc.add_paragraph("SKILLS")
    doc.add_paragraph("• Selenium WebDriver, Java, REST Assured")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

class TestParsing:
    def test_extracts_four_sections(self):
        buf = _create_test_docx()
        doc, sum_paras, exp_paras, proj_paras, skills_paras = parse_docx_sections(buf)
        assert len(sum_paras) == 1
        assert len(exp_paras) == 1
        assert len(proj_paras) == 1
        assert len(skills_paras) == 1
        assert "Experienced QA engineer" in get_texts(sum_paras)[0]
        assert "REST Assured" in get_texts(exp_paras)[0]

class TestMarkdownReplacement:
    def test_bold_markdown(self):
        doc = Document()
        para = doc.add_paragraph("• Original")
        replace_markdown_text(para, "This is **bold** text")
        # Should create 3 runs: "• This is ", "bold", " text"
        # However _parse logic prefix might inject early. 
        # Actually prefix goes into text before parse.
        # Let's just check the text content matches and at least one run is bold
        assert "This is bold text" in para.text
        assert any(r.bold for r in para.runs)

class TestAIParsing:
    def test_parses_4_sections(self):
        ai_text = """SUMMARY_TEXT:
New summary
EXPERIENCE_BULLETS:
• New exp
PROJECTS_BULLETS:
• New proj
SKILLS_TEXT:
New skills"""
        s, e, p, k = _parse_ai_4_sections(ai_text)
        assert s == ["New summary"]
        assert e == ["New exp"]
        assert p == ["New proj"]
        assert k == ["New skills"]

class TestPromptBuilder:
    def test_contains_all_texts(self):
        prompt = build_tailor_prompt("JD", ["SUM"], ["EXP"], ["PROJ"], ["SKILL"])
        assert "SUM" in prompt
        assert "EXP" in prompt
        assert "PROJ" in prompt
        assert "SKILL" in prompt

class TestEndToEnd:
    def test_full_pipeline(self):
        buf = _create_test_docx()
        doc, s, e, p, k = parse_docx_sections(buf)
        
        ai_response = """SUMMARY_TEXT:
**Experienced** QA
EXPERIENCE_BULLETS:
• Engineered REST
PROJECTS_BULLETS:
• Architected framework
SKILLS_TEXT:
Skills line"""

        count = apply_tailored_sections(s, e, p, k, ai_response)
        assert count == 4
        assert "**" not in s[0].text # asterisks removed
        assert any(r.bold for r in s[0].runs)
