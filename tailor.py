"""
tailor.py — Truthful Resume Tailor backend logic
=================================================
Handles:
  - Fetching 'Not Applied' jobs from Airtable
  - Parsing DOCX (Summary, Experience, Projects, Skills)
  - In-place DOCX run replacement with Markdown bolding
  - Gemini model cascade with 429 fallback
"""

import io
import re
import time
import datetime
import requests
import streamlit as st
import google.generativeai as genai
from docx import Document

from tailor_prompt import TAILOR_PROMPT

# ─── Constants ───────────────────────────────────────────────────────────────

AIRTABLE_BASE_ID = "appABPMwKgXkr8Rgn"
AIRTABLE_TABLE_NAME = "Applications"

MODEL_CASCADE = [
    ("gemini-3.1-flash-lite-preview", 15),
    ("gemini-3-flash-preview",         5),
    ("gemini-2.5-flash",               5),
]

DAILY_LIMIT = 540

SECTION_HEADERS = [
    "EXPERIENCE", "PROJECTS", "SKILLS", "EDUCATION",
    "PROFESSIONAL SUMMARY", "SUMMARY", "CERTIFICATIONS", "AWARDS",
]


# ─── Airtable: Fetch 'Not Applied' Jobs ─────────────────────────────────────

def fetch_not_applied_jobs(base_id: str, token: str) -> list[dict]:
    url = f"https://api.airtable.com/v0/{base_id}/{AIRTABLE_TABLE_NAME}"
    headers = {"Authorization": f"Bearer {token}"}
    params = {"filterByFormula": "{Status}='Not Applied'"}
    jobs = []
    offset = None
    while True:
        if offset:
            params["offset"] = offset
        try:
            resp = requests.get(url, headers=headers, params=params, timeout=30)
            resp.raise_for_status()
            data = resp.json()
        except Exception as e:
            st.error(f"Failed to fetch Airtable jobs: {e}")
            break
        for record in data.get("records", []):
            fields = record.get("fields", {})
            jobs.append({
                "record_id": record.get("id", ""),
                "company": fields.get("Company", "Unknown"),
                "role": fields.get("Role", "Unknown"),
                "jd_description": fields.get("JD Description", ""),
            })
        offset = data.get("offset")
        if not offset:
            break
    return jobs


# ─── DOCX Parsing ───────────────────────────────────────────────────────────

def _is_section_header(text: str) -> bool:
    stripped = text.strip().upper()
    return stripped in SECTION_HEADERS

def _is_bullet(paragraph) -> bool:
    text = paragraph.text.strip()
    if not text:
        return False
    if text.startswith("•") or text.startswith("-") or text.startswith("–"):
        return True
    style_name = (paragraph.style.name or "").lower()
    if "list" in style_name:
        return True
    return False

def parse_docx_sections(file) -> tuple:
    """
    Extract paragraphs for Summary, Experience, Projects, and Skills.
    """
    doc = Document(file)

    summary_paragraphs = []
    experience_paragraphs = []
    projects_paragraphs = []
    skills_paragraphs = []

    current_section = None

    for para in doc.paragraphs:
        text = para.text.strip()

        if _is_section_header(text):
            upper = text.upper()
            if upper in ["PROFESSIONAL SUMMARY", "SUMMARY"]:
                current_section = "summary"
            elif upper == "EXPERIENCE":
                current_section = "experience"
            elif upper == "PROJECTS":
                current_section = "projects"
            elif upper == "SKILLS":
                current_section = "skills"
            else:
                current_section = None
            continue

        if current_section == "summary" and text:
            summary_paragraphs.append(para)
        elif current_section == "experience" and _is_bullet(para):
            experience_paragraphs.append(para)
        elif current_section == "projects" and _is_bullet(para):
            projects_paragraphs.append(para)
        elif current_section == "skills" and text:
            skills_paragraphs.append(para)

    return doc, summary_paragraphs, experience_paragraphs, projects_paragraphs, skills_paragraphs

def get_texts(paragraphs: list) -> list[str]:
    """Extract clean text from paragraphs."""
    texts = []
    for para in paragraphs:
        text = para.text.strip()
        # Strip leading bullet character for clearer matching if it exists
        text = re.sub(r"^[•\-–]\s*", "", text)
        if text:
            texts.append(text)
    return texts


# ─── DOCX In-Place Replacement (Markdown Bolding) ───────────────────────────

def replace_markdown_text(paragraph, new_text: str):
    """
    Parse Markdown bold (**text**) and apply it to the DOCX paragraph.
    Clears existing runs and creates new ones, inheriting the original
    font attributes from the very first run to preserve formatting.
    """
    if not paragraph.runs:
        paragraph.text = new_text
        return

    # 1. Capture original font styling from the very first run
    original_font_name = paragraph.runs[0].font.name
    original_font_size = paragraph.runs[0].font.size

    # 2. Preserve bullet prefix natively if not provided by AI
    original_text = paragraph.text.strip()
    prefix = ""
    if original_text and original_text[0] in ("•", "-", "–") and not new_text.startswith(original_text[0]):
        prefix = original_text[0] + " "

    parsed_text = prefix + new_text

    # 3. Clear existing runs
    for run in paragraph.runs:
        run.text = ""

    # 4. Split by ** to capture bold chunks
    parts = re.split(r'(\*\*.*?\*\*)', parsed_text)

    for part in parts:
        if not part:
            continue
        is_bold = False
        if part.startswith("**") and part.endswith("**"):
            is_bold = True
            part = part[2:-2] # remove the ** markers
        
        # Translate explicit literal \n strings to actual line breaks
        lines = part.replace("\\n", "\n").split("\n")
        
        for i, line_text in enumerate(lines):
            if line_text:
                new_run = paragraph.add_run(line_text)
                new_run.bold = is_bold
                if original_font_name is not None:
                    new_run.font.name = original_font_name
                if original_font_size is not None:
                    new_run.font.size = original_font_size
            
            # Add a vertical break if not the last line chunk
            if i < len(lines) - 1:
                paragraph.add_run().add_break()


def apply_tailored_sections(
    summary_paras, exp_paras, proj_paras, skills_paras, ai_response: str
):
    """Parse the 4 sections from AI and apply them to the DOCX."""
    summary_res, exp_res, proj_res, skills_res = _parse_ai_4_sections(ai_response)

    count = 0
    for i, para in enumerate(summary_paras):
        if i < len(summary_res) and summary_res[i] != "(none)":
            text = summary_res[i]
            # Ensure the summary paragraph natively has breathing room (spacing above/below)
            if not text.startswith("\n") and not text.startswith("\\n"):
                text = "\n" + text
            if not text.endswith("\n") and not text.endswith("\\n"):
                text = text + "\n"
            replace_markdown_text(para, text)
            count += 1

    for i, para in enumerate(exp_paras):
        if i < len(exp_res) and exp_res[i] != "(none)":
            replace_markdown_text(para, exp_res[i])
            count += 1

    for i, para in enumerate(proj_paras):
        if i < len(proj_res) and proj_res[i] != "(none)":
            replace_markdown_text(para, proj_res[i])
            count += 1

    for i, para in enumerate(skills_paras):
        if i < len(skills_res) and skills_res[i] != "(none)":
            replace_markdown_text(para, skills_res[i])
            count += 1

    return count


def _parse_ai_4_sections(ai_text: str) -> tuple[list[str], list[str], list[str], list[str]]:
    text = ai_text.replace("\r\n", "\n")
    
    # Simple state machine to parse the 4 sections
    sections = {
        "SUMMARY_TEXT": [],
        "EXPERIENCE_BULLETS": [],
        "PROJECTS_BULLETS": [],
        "SKILLS_TEXT": []
    }
    
    current_key = None
    
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
            
        is_header = False
        for key in sections.keys():
            if line.upper().startswith(key):
                current_key = key
                is_header = True
                break
                
        if is_header:
            continue
            
        if current_key and line != "(none)":
            # Remove leading bullet if the AI provided one to avoid double bulleting, 
            # replace_markdown_text adds it back if the original para had it.
            clean_line = re.sub(r"^[•\-–]\s*", "", line)
            sections[current_key].append(clean_line)
            
    return (
        sections["SUMMARY_TEXT"],
        sections["EXPERIENCE_BULLETS"],
        sections["PROJECTS_BULLETS"],
        sections["SKILLS_TEXT"]
    )


# ─── Gemini Model Cascade ───────────────────────────────────────────────────

def _init_rate_state():
    if "tailor_api_calls" not in st.session_state:
        st.session_state["tailor_api_calls"] = 0
        st.session_state["tailor_api_date"] = datetime.date.today().isoformat()
    # Reset counter at midnight
    if st.session_state["tailor_api_date"] != datetime.date.today().isoformat():
        st.session_state["tailor_api_calls"] = 0
        st.session_state["tailor_api_date"] = datetime.date.today().isoformat()

def is_rate_limited() -> bool:
    _init_rate_state()
    return st.session_state["tailor_api_calls"] >= DAILY_LIMIT


def generate_with_fallback(prompt: str, api_key: str) -> str | None:
    _init_rate_state()
    if is_rate_limited():
        st.error(f"⚠️ Daily Limit Reached ({DAILY_LIMIT}). Try tomorrow.")
        return None

    genai.configure(api_key=api_key)

    for model_name, _rpm in MODEL_CASCADE:
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(prompt)
            st.session_state["tailor_api_calls"] += 1
            return response.text
        except Exception as e:
            error_str = str(e)
            if "429" in error_str or "RESOURCE_EXHAUSTED" in error_str:
                st.toast(f"⚠️ {model_name} rate-limited, falling back...", icon="🔄")
                time.sleep(5)
                continue
            else:
                st.error(f"❌ Error with {model_name}: {e}")
                return None
    st.error("❌ All models exhausted. Please wait a few minutes and retry.")
    return None


# ─── Prompt Builder ──────────────────────────────────────────────────────────

def build_tailor_prompt(
    jd_text: str,
    summary_texts: list[str],
    experience_texts: list[str],
    projects_texts: list[str],
    skills_texts: list[str]
) -> str:
    sum_fmt  = "\n".join(summary_texts) if summary_texts else "(none)"
    exp_fmt  = "\n".join(f"• {b}" for b in experience_texts) if experience_texts else "(none)"
    proj_fmt = "\n".join(f"• {b}" for b in projects_texts) if projects_texts else "(none)"
    sk_fmt   = "\n".join(skills_texts) if skills_texts else "(none)"

    return TAILOR_PROMPT.format(
        jd_text=jd_text,
        summary_text=sum_fmt,
        experience_bullets=exp_fmt,
        projects_bullets=proj_fmt,
        skills_text=sk_fmt,
    )


# ─── Save Modified DOCX to BytesIO ──────────────────────────────────────────

def save_doc_to_bytes(doc) -> bytes:
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
